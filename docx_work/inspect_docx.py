from __future__ import annotations

import sys
from collections import Counter
from pathlib import Path

from docx import Document


def twips_to_in(value):
    return None if value is None else round(value.twips / 1440, 3)


def fmt_para(p):
    pf = p.paragraph_format
    rf = p.runs[0].font if p.runs else None
    text = p.text.strip().replace("\n", " ")
    return {
        "style": p.style.name if p.style else "",
        "align": p.alignment,
        "left": twips_to_in(pf.left_indent),
        "right": twips_to_in(pf.right_indent),
        "first": twips_to_in(pf.first_line_indent),
        "before": None if pf.space_before is None else round(pf.space_before.pt, 2),
        "after": None if pf.space_after is None else round(pf.space_after.pt, 2),
        "line": pf.line_spacing,
        "font": None if rf is None else rf.name,
        "size": None if rf is None or rf.size is None else round(rf.size.pt, 2),
        "bold": None if rf is None else rf.bold,
        "italic": None if rf is None else rf.italic,
        "sample": text[:90],
    }


def inspect(path: Path):
    doc = Document(path)
    print(f"FILE: {path}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
    for i, sec in enumerate(doc.sections):
        print(
            "SECTION",
            i,
            f"page={twips_to_in(sec.page_width)}x{twips_to_in(sec.page_height)}",
            f"margins T/R/B/L={twips_to_in(sec.top_margin)}/{twips_to_in(sec.right_margin)}/{twips_to_in(sec.bottom_margin)}/{twips_to_in(sec.left_margin)}",
            f"header={twips_to_in(sec.header_distance)} footer={twips_to_in(sec.footer_distance)}",
        )
    styles = Counter(p.style.name if p.style else "" for p in doc.paragraphs)
    print("TOP STYLES:", styles.most_common(12))
    print("FIRST 60 NONEMPTY PARAGRAPHS:")
    shown = 0
    for idx, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        print(idx, fmt_para(p))
        shown += 1
        if shown >= 60:
            break


if __name__ == "__main__":
    for arg in sys.argv[1:]:
        inspect(Path(arg))
        print()
