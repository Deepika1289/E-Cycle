from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TITLE_WORDS = {
    "DECLARATION",
    "CERTIFICATE",
    "ACKNOWLEDGMENT",
    "ACKNOWLEDGEMENTS",
    "CONTENTS",
    "TABLE OF CONTENTS",
    "LIST OF ABBREVIATIONS",
    "LIST OF TABLES",
    "ABSTRACT",
    "ABSTRACT:-",
    "INTRODUCTION",
}


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip()).upper()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size: float | None = None, bold=None, italic=None) -> None:
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic


def format_para(p, *, size=12, bold=None, italic=None, align=None, first=0.0,
                left=0.0, right=0.0, before=0, after=0, line=1.5) -> None:
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.left_indent = Inches(left)
    pf.right_indent = Inches(right)
    pf.first_line_indent = Inches(first)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    for run in p.runs:
        set_run_font(run, size=size, bold=bold, italic=italic)


def has_page_break_before(p) -> bool:
    xml = p._p.xml
    return 'w:type="page"' in xml or "lastRenderedPageBreak" in xml


def add_page_break_before(p) -> None:
    if has_page_break_before(p):
        return
    previous = p.insert_paragraph_before()
    previous.add_run().add_break(WD_BREAK.PAGE)


def is_probable_heading(text: str) -> bool:
    c = clean(text).rstrip(":")
    raw = re.sub(r"\s+", " ", text.strip()).rstrip(":")
    if not c:
        return False
    if c in TITLE_WORDS:
        return True
    if re.match(r"^CHAPTER\s*[- ]?\s*[0-9IVX]+$", c):
        return True
    if re.match(r"^CHAPTER\s*[- ]?\s*[0-9IVX]+[:. -].+", c):
        return True
    if re.match(r"^(AIM|AIMS|OBJECTIVE|OBJECTIVES|MATERIALS AND METHODS|RESULTS?|DISCUSSION|CONCLUSION|SUMMARY|BIBLIOGRAPHY|REFERENCES)\s*$", c):
        return True
    return len(raw) <= 55 and raw == raw.upper() and any(ch.isalpha() for ch in raw)


def style_document(path_in: Path, path_out: Path) -> None:
    shutil.copyfile(path_in, path_out)
    doc = Document(path_out)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.page_width = Inches(8.271)
        section.page_height = Inches(11.694)
        section.top_margin = Inches(0.681)
        section.bottom_margin = Inches(0.556)
        section.left_margin = Inches(0.492)
        section.right_margin = Inches(0.492)
        section.header_distance = Inches(0.499)
        section.footer_distance = Inches(0.426)

    page_start_markers = {
        "DECLARATION",
        "CERTIFICATE",
        "ACKNOWLEDGMENT",
        "ACKNOWLEDGEMENTS",
        "CONTENTS",
        "TABLE OF CONTENTS",
        "LIST OF ABBREVIATIONS",
        "LIST OF TABLES",
        "ABSTRACT",
        "ABSTRACT:-",
        "DISCUSSION",
        "CONCLUSION",
        "REFERENCES",
        "BIBLIOGRAPHY",
    }

    for p in doc.paragraphs:
        marker = clean(p.text).rstrip(":")
        if not marker:
            continue
        if marker in page_start_markers or re.match(r"^CHAPTER\s*[- ]?\s*[0-9IVX]+$", marker):
            add_page_break_before(p)

    nonempty_seen = 0
    after_abstract = False
    in_contents = False
    in_list_section = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            continue

        marker = clean(text)
        nonempty_seen += 1

        if marker in {"CONTENTS", "TABLE OF CONTENTS"}:
            in_contents = True
            in_list_section = False
        elif marker in {"LIST OF ABBREVIATIONS", "LIST OF TABLES"}:
            in_contents = False
            in_list_section = True
        elif marker.startswith("ABSTRACT"):
            in_contents = False
            in_list_section = False
        elif re.match(r"^CHAPTER\s*[- ]?\s*[0-9IVX]+$", marker):
            in_contents = False
            in_list_section = False

        if nonempty_seen <= 17:
            # Cover/front matter: template uses centered title-page treatment.
            size = 14
            bold = None
            italic = None
            before = 0
            if nonempty_seen <= 3:
                size, bold = 20, True
            elif "DEGREE" in marker or "BACHELOR" in marker:
                size, bold = 18, True
                before = 12
            elif marker in {"BY", "UNDER THE GUIDANCE OF"}:
                size, italic = 14, True
                before = 8
            elif "NUKALA" in marker or "REGD" in marker:
                size, bold = 16, True
            elif "DEPARTMENT" in marker or "SCHOOL" in marker or "CENTURION" in marker:
                size, bold = 14, True
            format_para(p, size=size, bold=bold, italic=italic, align=WD_ALIGN_PARAGRAPH.CENTER, before=before, line=1.15)
            continue

        if in_contents and marker not in {"CONTENTS", "TABLE OF CONTENTS"}:
            format_para(p, size=12, bold=False, italic=False, align=None, first=0, left=0.1, right=0.1, before=0, after=0, line=1.15)
            continue

        if in_list_section and marker not in {"LIST OF ABBREVIATIONS", "LIST OF TABLES"}:
            format_para(p, size=12, bold=False, italic=False, align=None, first=0, left=0.1, right=0.1, before=0, after=0, line=1.15)
            continue

        if re.match(r"^CHAPTER\s*[- ]?\s*[0-9IVX]+$", marker):
            after_abstract = False
            format_para(p, size=48, bold=True, italic=None, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=1.0)
            continue

        if marker.startswith(("PLACE", "DATE")) or "SIGNATURE" in marker or "REGD NO" in marker:
            format_para(p, size=12, bold=None, italic=None, align=None, first=0, left=0.1, before=7, line=1.15)
            continue

        if is_probable_heading(text):
            size = 14
            before = 16 if marker in TITLE_WORDS or marker.startswith("CHAPTER") else 8
            format_para(p, size=size, bold=True, italic=None, align=WD_ALIGN_PARAGRAPH.CENTER, before=before, after=0, line=1.15)
            if marker.startswith("ABSTRACT"):
                after_abstract = True
            elif marker.startswith("CHAPTER"):
                after_abstract = False
            continue

        body_size = 14 if after_abstract and nonempty_seen < 90 else 12
        format_para(
            p,
            size=body_size,
            bold=False,
            italic=False,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first=0.5,
            left=0.1,
            right=0.1,
            before=0,
            after=0,
            line=1.5,
        )

    for table in doc.tables:
        table.autofit = True
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if r_idx == 0:
                    set_cell_shading(cell, "EDEDED")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    for run in p.runs:
                        set_run_font(run, size=10 if r_idx else 10.5, bold=True if r_idx == 0 else None)

    doc.save(path_out)


if __name__ == "__main__":
    style_document(Path(sys.argv[1]), Path(sys.argv[2]))
