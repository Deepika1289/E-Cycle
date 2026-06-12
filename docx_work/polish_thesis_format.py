from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


MAJOR_TITLES = {
    "DECLARATION",
    "CERTIFICATE",
    "ACKNOWLEDGMENT",
    "ACKNOWLEDGEMENTS",
    "CONTENTS",
    "TABLE OF CONTENTS",
    "LIST OF ABBREVIATIONS",
    "LIST OF TABLES",
    "ABSTRACT",
    "DISCUSSION",
    "CONCLUSION",
    "REFERENCES",
    "BIBLIOGRAPHY",
}


def norm(text: str) -> str:
    return re.sub(r"[\s:.-]+$", "", re.sub(r"\s+", " ", text.strip()).upper())


def set_run(run, *, size=12, bold=False, italic=False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def apply_para(p, *, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
               left=0, right=0, first=0, before=0, after=0, line=1.5) -> None:
    p.alignment = align
    pf = p.paragraph_format
    pf.left_indent = Inches(left)
    pf.right_indent = Inches(right)
    pf.first_line_indent = Inches(first)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    for run in p.runs:
        set_run(run, size=size, bold=bold, italic=italic)


def add_break_before(p) -> None:
    if 'w:type="page"' in p._p.xml:
        return
    pb = p.insert_paragraph_before()
    pb.add_run().add_break(WD_BREAK.PAGE)


def set_cell_margins(cell, value=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def is_chapter(text: str) -> bool:
    return bool(re.match(r"^CHAPTER\s*[- ]?\s*[0-9IVX]+$", norm(text)))


def is_major_title(text: str) -> bool:
    t = norm(text)
    return t in MAJOR_TITLES or is_chapter(text)


def is_short_standalone_heading(text: str) -> bool:
    raw = re.sub(r"\s+", " ", text.strip()).rstrip(":")
    up = raw.upper()
    if not raw or ":" in raw or len(raw) > 70:
        return False
    if up in MAJOR_TITLES or is_chapter(raw):
        return True
    return raw == up and any(ch.isalpha() for ch in raw)


def nonempty_paragraphs(doc: Document):
    return [p for p in doc.paragraphs if p.text.strip()]


def polish(input_path: Path, output_path: Path) -> None:
    shutil.copyfile(input_path, output_path)
    doc = Document(output_path)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    nonempty = nonempty_paragraphs(doc)
    page_break_targets = set()
    for i, p in enumerate(nonempty):
        t = norm(p.text)
        prev = norm(nonempty[i - 1].text) if i else ""
        if i == 0:
            continue
        if is_chapter(p.text) or t in MAJOR_TITLES:
            page_break_targets.add(id(p))
        if t.startswith("CENTURION UNIVERSITY OF TECHNOLOGY") and "VIZIANAGARAM" in t and prev.startswith("DATE"):
            page_break_targets.add(id(p))

    for p in doc.paragraphs:
        if id(p) in page_break_targets:
            add_break_before(p)

    nonempty = nonempty_paragraphs(doc)
    index_by_id = {id(p): i for i, p in enumerate(nonempty)}
    cover_limit = min(17, len(nonempty))
    in_contents = False
    in_lists = False
    in_abstract = False
    current_chapter = False

    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            continue

        marker = norm(text)
        nonempty_index = index_by_id.get(id(p), 9999)

        if nonempty_index < cover_limit:
            if nonempty_index <= 2:
                apply_para(p, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4, line=1.15)
            elif "DEGREE" in marker:
                apply_para(p, size=13, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=4, line=1.15)
            elif "BACHELOR" in marker:
                apply_para(p, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=12, line=1.15)
            elif marker in {"BY", "UNDER THE GUIDANCE OF"}:
                apply_para(p, size=13, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=2, line=1.15)
            elif "NUKALA" in marker or "REGD" in marker:
                apply_para(p, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2, line=1.15)
            else:
                apply_para(p, size=13, bold=("CENTURION" in marker or "DEPARTMENT" in marker or "SCHOOL" in marker), align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2, line=1.15)
            continue

        if marker in {"CONTENTS", "TABLE OF CONTENTS"}:
            in_contents, in_lists, in_abstract, current_chapter = True, False, False, False
        elif marker in {"LIST OF ABBREVIATIONS", "LIST OF TABLES"}:
            in_contents, in_lists, in_abstract, current_chapter = False, True, False, False
        elif marker == "ABSTRACT":
            in_contents, in_lists, in_abstract, current_chapter = False, False, True, False
        elif is_chapter(text):
            in_contents, in_lists, in_abstract, current_chapter = False, False, False, True
        elif marker in {"REFERENCES", "BIBLIOGRAPHY"}:
            in_contents, in_lists, in_abstract, current_chapter = False, False, False, False

        if is_chapter(text):
            apply_para(p, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10, line=1.15)
        elif marker in MAJOR_TITLES:
            apply_para(p, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=14, line=1.15)
        elif in_contents or in_lists:
            apply_para(p, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, left=0.15, first=0, before=0, after=4, line=1.15)
        elif is_short_standalone_heading(text):
            apply_para(p, size=13, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT if current_chapter else WD_ALIGN_PARAGRAPH.CENTER, before=12, after=6, line=1.15)
        elif marker.startswith(("PLACE", "DATE")) or "SIGNATURE" in marker:
            apply_para(p, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, left=0, first=0, before=12, after=0, line=1.15)
        else:
            apply_para(p, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left=0, right=0, first=0.5, before=0, after=6 if in_abstract else 3, line=1.5)

    for table in doc.tables:
        table.autofit = True
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, 90)
                if r_idx == 0:
                    set_cell_shading(cell, "EDEDED")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    for run in p.runs:
                        set_run(run, size=10.5 if r_idx == 0 else 10, bold=(r_idx == 0), italic=False)

    doc.save(output_path)


if __name__ == "__main__":
    polish(Path(sys.argv[1]), Path(sys.argv[2]))
